import sys
import os

# --- RIGOROUS ENVIRONMENT PROTOCOL ---
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(os.path.dirname(current_dir))
venv_site_packages = os.path.join(project_root, 'venv', 'lib', 'python3.9', 'site-packages')

if os.path.exists(venv_site_packages):
    sys.path.insert(0, venv_site_packages)

try:
    from docx import Document
except ImportError:
    Document = None
# -------------------------------------

def extract(file_path):
    """
    Exhaustively explores .docx technical documents.
    Zero truncation: identifies tables, images, and core properties.
    """
    if not Document:
        return "Error: python-docx library not found in venv."

    try:
        doc = Document(file_path)
        prop = doc.core_properties
        
        # 1. Structural Audit
        table_count = len(doc.tables)
        image_count = len(doc.inline_shapes)
        para_count = len(doc.paragraphs)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)

        # 2. Extract Timing
        created = prop.created.strftime('%Y-%m-%d') if prop.created else "Unknown"
        modified = prop.modified.strftime('%Y-%m-%d') if prop.modified else "Unknown"

        output = [
            f"Type: Office Open XML Document (.docx)",
            f"Subject: {prop.subject if prop.subject else 'Technical Report'}",
            f"Author: {prop.author if prop.author else 'Unknown'}",
            f"Timeline: Created {created} | Last Mod {modified}",
            "Structural Density:",
            f"  - Content: {para_count} Paragraphs | ~{word_count} Words",
            f"  - Assets:  {table_count} Tables | {image_count} Images/Graphics",
            f"Keywords: {prop.keywords if prop.keywords else 'None'}"
        ]

        return "\n".join(output)

    except Exception as e:
        return f"DOCX Extraction Error: {str(e)}"