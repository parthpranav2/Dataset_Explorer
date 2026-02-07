import sys
import os
import datetime

# --- RIGOROUS ENVIRONMENT PROTOCOL ---
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(os.path.dirname(current_dir))
venv_site_packages = os.path.join(project_root, 'venv', 'lib', 'python3.9', 'site-packages')

if os.path.exists(venv_site_packages):
    sys.path.insert(0, venv_site_packages)

try:
    from docx import Document
except ImportError:
    Document = None
# -------------------------------------

def extract(file_path):
    """
    Exhaustively explores .doc/.docx technical reports.
    Zero truncation: provides metadata, word counts, and structural stats.
    """
    try:
        file_size = os.path.getsize(file_path)
        stats = {
            "Author": "Unknown",
            "Created": "Unknown",
            "Pages": "N/A",
            "Words": "N/A"
        }

        # Handling modern .docx
        if file_path.lower().endswith('.docx') and Document:
            doc = Document(file_path)
            prop = doc.core_properties
            stats["Author"] = prop.author if prop.author else "Unknown"
            stats["Created"] = prop.created.strftime('%Y-%m-%d') if prop.created else "Unknown"
            
            # Count paragraphs and rough word count
            paragraphs = len(doc.paragraphs)
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            stats["Words"] = word_count
            
            output = [
                f"Type: Word Document (.docx)",
                f"Document Title: {prop.title if prop.title else os.path.basename(file_path)}",
                f"Author: {stats['Author']}",
                f"Created On: {stats['Created']}",
                f"Structure: {paragraphs} Paragraphs | ~{stats['Words']} Words",
                f"File Size: {file_size / 1024:.2f} KB"
            ]
        
        # Handling legacy .doc (Basic Stat Extraction)
        else:
            mod_time = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d')
            output = [
                f"Type: Legacy Word Document (.doc)",
                f"Note: Deep parsing requires .docx format.",
                f"Last Modified: {mod_time}",
                f"File Size: {file_size / 1024:.2f} KB"
            ]

        return "\n".join(output)

    except Exception as e:
        return f"DOC Extraction Error: {str(e)}"